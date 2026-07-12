from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX_PATH = Path(__file__).with_name("The_Neighbourhood_Launch_Execution_Plan_15_Aug_2026.docx")
TEMP_PATH = DOCX_PATH.with_suffix(".tmp.docx")
BLUE = "1F4D78"
INK = "1D1D1F"


def set_font(run, color=INK, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    return run


def set_label(paragraph, label, value):
    paragraph.clear()
    lead = paragraph.add_run(f"{label}: ")
    set_font(lead, color=BLUE, bold=True)
    body = paragraph.add_run(value)
    set_font(body)


def set_cell(cell, text):
    while len(cell.paragraphs) > 1:
        element = cell.paragraphs[-1]._element
        element.getparent().remove(element)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(8.65)
    run.font.color.rgb = RGBColor.from_string(INK)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def find_paragraph(doc, starts_with):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(starts_with):
            return paragraph
    raise ValueError(f"Paragraph not found: {starts_with}")


def replace_label(doc, label, value):
    set_label(find_paragraph(doc, f"{label}:"), label, value)


def remove_milestone_spec(doc):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.startswith("3.6. Milestone guide"))
    end = next(i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.startswith("3.7. Safe AI copilot"))
    for paragraph in paragraphs[start:end]:
        paragraph._element.getparent().remove(paragraph._element)


def revise_mvp_table(doc):
    table = doc.tables[2]
    today = table.rows[4]
    values = [
        "4. Today: daily guidance and milestone context",
        "This is the daily reason to return. It gives the parent one achievable moment and just enough context about what may be unfolding.",
        "One featured activity plus three visible alternatives, clear steps and reason, optional multiple commitments, and a compact 'What may be unfolding' section inside Today with two or three age-relevant cues.",
        "A separate milestone destination, checklist score, comparisons, diagnosis, automated alerts, AI-generated activities, time-of-day prediction, streaks, badges, and forced completion.",
        "Child age, curated activity and milestone content, daily-plan service, optional observation table.",
        "M: 5-6 days",
    ]
    for cell, value in zip(today.cells, values):
        set_cell(cell, value)

    milestone_row = table.rows[6]
    milestone_row._tr.getparent().remove(milestone_row._tr)

    renumber = {
        "7. Safe AI copilot": "6. Safe AI copilot",
        "8. Lab and counselling booking request": "7. Lab and counselling booking request",
        "9. Feedback and product measurement": "8. Feedback and product measurement",
        "10. Content and operations workflow": "9. Content and operations workflow",
        "11. Landing and invite path": "10. Landing and invite path",
    }
    for row in table.rows:
        if row.cells[0].text in renumber:
            set_cell(row.cells[0], renumber[row.cells[0].text])

    set_cell(
        doc.tables[3].cell(0, 0),
        "Capacity check\nThis is roughly 34-40 focused engineering days after content is ready. The 15 August date is realistic with two engineers working in parallel, a named content owner, and a named bookings owner. With one engineer, ship Today with its milestone context, copilot, account/onboarding, analytics, and a contact/request form; move the richer milestone note capability and slot picker to Version 1.1.",
    )


def revise_flows(doc):
    set_text(find_paragraph(doc, "Open saved link or type URL"), "Open saved link or type URL\n  |\n  v\nRestore session\n  |\n  +-- Yesterday's optional reflection, if relevant\n  |\n  v\nToday's four activities and a small 'what may be unfolding' cue\n  |\n  v\nCopilot when needed")
    replace_label(doc, "What the parent sees", "The page opens on Today, never on a dashboard. If yesterday's plan was not reflected on, a compact card appears first. It can be completed, deferred, or quietly dismissed. The activity choices and a small 'what may be unfolding' cue stay together on Today.")
    replace_label(doc, "After each click", "Tapping an activity expands its steps in place. 'I'll make room for this' toggles a voluntary intention. 'What may be unfolding' expands in place; it is not a separate milestone page. 'Ask the copilot' opens chat in context.")

    heading = find_paragraph(doc, "2.4 Today -> Milestone")
    set_text(heading, "2.4 Today -> What may be unfolding -> Detail -> Private note")
    set_text(find_paragraph(doc, "Today\n  |\n  v\nMilestones"), "Today\n  |\n  v\nExpand 'What may be unfolding'\n  |\n  v\nAge-relevant cue\n  |\n  v\nDetail\n  |\n  v\nOptional 'I noticed this'\n  |\n  v\nBack to the same Today page")
    replace_label(doc, "What the parent sees", "A compact, expandable part of Today, not a separate destination. It starts with the age range and a reminder that children develop at different paces. Each cue says what a parent may notice and one simple thing to try.")
    replace_label(doc, "After each click", "Opening a cue reveals its detail in place. 'I noticed this' saves a private date stamp and changes the cue to 'Noted by you'. 'Undo' removes it. There is no list of untouched items labelled incomplete.")
    replace_label(doc, "Data that changes", "milestone_observations: child_id, milestone_id, observed_at, optional note. milestone_cue_view and milestone_noted analytics events.")
    replace_label(doc, "Edge cases", "No content for the stage: hide the cue rather than showing an empty milestone section. A parent flags a concern: do not interpret it clinically; invite them to ask the copilot or request a conversation. Duplicate note: toggle, do not create duplicates.")


def revise_today_spec(doc):
    set_text(find_paragraph(doc, "3.4. Today daily guidance"), "3.4. Today daily guidance and milestone context")
    replace_label(doc, "Purpose", "Give the parent one clear, useful action for this child today, with a little developmental context only when they want it.")
    replace_label(doc, "User story", "As a parent, I want one small, realistic way to be with my child today and a calm sense of what may be unfolding at this age.")
    replace_label(doc, "Screen", "Today page: brief greeting, optional reflection card, one featured activity, three alternatives, a compact expandable 'What may be unfolding' block, copilot, sessions, and profile.")
    replace_label(doc, "Inputs", "Child age stage, approved activity and milestone content, daily rotation, recent optional observation, voluntary intentions, and optional private milestone notes.")
    replace_label(doc, "Outputs", "One deterministic four-activity daily plan, activity detail, age-relevant milestone cues, optional intentions and notes, and activity/milestone analytics.")
    replace_label(doc, "States", "Loading, ready, expanded activity, expanded milestone cue, committed, cue noted, offline cached, and no-content fallback.")
    replace_label(doc, "Buttons", "Open activity, close detail, 'I'll make room for this', remove from today, open 'What may be unfolding', open cue detail, 'I noticed this', undo, ask the copilot, and view sessions.")
    replace_label(doc, "Actions", "Generate the daily plan server-side or with a shared deterministic rule. Keep four activities stable for the day. Render two or three approved age-relevant cues inside Today. Let parents commit to zero, one, or many and optionally note a cue. Selecting any card or cue must not reorder the layout. Use no emoji decoration, streak, badge, score, or 'missed' wording.")
    replace_label(doc, "Business rules", "Activities and milestone cues are selected from approved age-stage content only. Today is the only launch surface for milestone context. A commitment or note is optional and not a promise or proof of development. Never calculate progress, show achievement percentages, call anything missed, compare children, or infer delay. The page must remain usable at 320px wide and use at least 44px touch targets.")
    replace_label(doc, "Empty state", "If activity content is unavailable, show a single evergreen prompt with retry. If milestone content is unavailable, hide that small section. Never show an empty card grid or a separate empty milestone screen.")
    replace_label(doc, "Loading state", "Use a page skeleton that preserves the headline, activity-card heights, and the compact milestone cue area; do not animate aggressively.")
    replace_label(doc, "Error state", "Keep a cached plan visible if possible. Otherwise show a retry and a short support route. Preserve a private milestone note locally for retry.")
    replace_label(doc, "Success state", "The parent leaves knowing one possible next moment, with any selected activities clearly marked as part of today and any private observation quietly noted.")
    replace_label(doc, "Future improvements", "Richer milestone timeline or export for a clinician at a parent's request, context from proven memory, parent-selected energy level, saved activities, and accessible audio guidance. Do not build predictive AI recommendations before the content loop is trusted.")


def renumber_feature_specs(doc):
    changes = {
        "3.7. Safe AI copilot": "3.6. Safe AI copilot",
        "3.8. Lab and counselling booking requests": "3.7. Lab and counselling booking requests",
        "3.9. Feedback, analytics, and operational review": "3.8. Feedback, analytics, and operational review",
        "3.10. Content and operations workflow": "3.9. Content and operations workflow",
        "3.11. Landing, waitlist, and invitation path": "3.10. Landing, waitlist, and invitation path",
    }
    for old, new in changes.items():
        set_text(find_paragraph(doc, old), new)


def revise_schedule(doc):
    table = doc.tables[5]
    rows = [
        (
            "11-13 Jul\nFoundation",
            "Lock the schema and access rules. Set up Supabase Auth, profiles, child records, consent, error tracking, event plumbing, content tables/import, and the first approved content set. Keep the landing/waitlist working.",
            "Every later screen needs a secure identity, one source of truth, and approved content. This is the shortest viable foundation window.",
            "A parent can sign in on mobile, create one child, return on a second device, and the team can see the onboarding funnel.",
        ),
        (
            "14-17 Jul\nToday + milestone context",
            "Build Today on real data: four stable activities, activity detail, multiple optional commitments, next-day reflection, quiet stale state, profile editing, and an in-page 'What may be unfolding' block with optional private notes. Migrate the useful current local-state behavior carefully.",
            "Today is the daily value loop. Milestone context belongs inside it, so this is one cohesive build rather than a second destination.",
            "Five internal testers can complete onboarding, use Today for two simulated days, expand an age-relevant cue, optionally note it, edit their profile, and see the same state after sign-out/sign-in.",
        ),
        (
            "18-24 Jul\nCopilot + booking",
            "Build the safe copilot and booking request flow. Add reviewable knowledge-base retrieval, topic/source labels, medical escalation, answer feedback, session content, pending booking request, and operations queue. Run the full mobile-browser QA script.",
            "These two flows deliver the distinct value in the shared plan: useful real questions answered safely and a bridge to the physical Neighbourhood.",
            "Five parents complete an unassisted session. Copilot feedback works. A medical question routes safely. A booking request reaches the named operations owner.",
        ),
        (
            "25-31 Jul\nParent testing + 50",
            "Run cold-start testing with 10 new parents, then a controlled push to 50. Fix the three highest-frequency issues daily. Review every copilot miss, complete the content gaps, measure onboarding time and drop-off, and test slow 4G plus small Android screens.",
            "The product should now earn evidence, not more features. This is where unclear labels, slow paths, and false assumptions get removed.",
            "Onboarding is under 3 minutes without help; each key flow has been observed; no P0/P1 defects remain; copilot helpful rate is moving toward 80%; operations handles booking requests.",
        ),
        (
            "1-7 Aug\nHarden + 100",
            "Harden, launch in cohorts, and aim for 100 active parents. Monitor analytics each morning. Fix one proven friction point per day. Freeze features except for defects, content fixes, safety, and conversion blockers. Complete handoff notes and incident contacts.",
            "At this point acquisition and retention depend on reliability. Feature churn creates risk and makes the data impossible to read.",
            "100 active parents is the goal, but do not fake it with accounts. The app has a written support path, known issues, rollback plan, privacy review, and a clear owner for content, bookings, and technical incidents.",
        ),
        (
            "8-12 Aug\nLaunch freeze",
            "Run final mobile-browser QA, production smoke tests, support drills, and a final content/safety check. Invite the final planned cohort. No new product features.",
            "A quiet launch freeze protects stability and gives the team time to rehearse real support and operations.",
            "All launch-critical paths pass on a real Android and iPhone browser. The content, booking, and technical owners have tested their operating steps.",
        ),
        (
            "13-15 Aug\nLaunch",
            "Release candidate, controlled launch, daily operating rhythm, issue triage, and launch review. No native app work.",
            "A stable web launch is more valuable than a rushed second platform. The team needs time to respond to real parents.",
            "The live cohort can use every launch-critical flow, owners are reachable, and only defects, safety issues, or content corrections enter the release.",
        ),
    ]

    while len(table.rows) - 1 < len(rows):
        new_row = table.add_row()
        cant_split(new_row)

    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def revise_open_question(doc):
    table = doc.tables[11]
    for row in table.rows[1:]:
        if row.cells[0].text == "Milestone concern copy":
            set_cell(row.cells[1], "Poorly worded guidance inside Today can create anxiety or hide an important concern.")
            set_cell(row.cells[3], "Keep concern guidance inside expanded Today detail, never in a summary score. Have it clinically reviewed and link to human support, not AI diagnosis.")


def main():
    doc = Document(DOCX_PATH)

    revise_mvp_table(doc)
    revise_flows(doc)
    revise_today_spec(doc)
    remove_milestone_spec(doc)
    renumber_feature_specs(doc)
    revise_schedule(doc)
    revise_open_question(doc)

    doc.save(TEMP_PATH)
    TEMP_PATH.replace(DOCX_PATH)


if __name__ == "__main__":
    main()
